# -*- coding:utf-8 -*-

import os

import random
import sys
import datetime
from docx import Document
from flask import render_template, redirect, url_for, request, send_file, send_from_directory, jsonify
from werkzeug.utils import secure_filename
import xlrd

from app import db
from app.models import Question, QuestionType, KnowledgePoint, Subject, Chapter, User
from app.ques import ques
from app.ques.forms import OrganizeQuestionForm, EditQuestionForm, AddQuestionForm, AnswerForm, SubjectForm, ChapterForm

reload(sys)
sys.setdefaultencoding('utf8')


@ques.route('/knp/<string:kn>', methods=['GET', 'POST'])
def kn_point(kn):
    kn_id = KnowledgePoint.query.filter_by(name=kn).first().id
    results = Question.query.filter_by(kn_point_id=kn_id).all()
    return render_template('search_results.html', results=results, count=len(results))


@ques.route('/edit_ques/<int:qid>', methods=['GET', 'POST'])
def edit_ques(qid):
    form = EditQuestionForm()
    q = Question.query.filter_by(id=qid).first()
    qt = QuestionType.query.filter_by(id=q.question_type_id).first()
    form.context.data = q.context
    form.answer.data = q.answer
    form.kn_point.data = KnowledgePoint.query.filter_by(id=q.kn_point_id).first().name
    if qt.name == "选择题":
        form.options.data = q.options

    if form.validate_on_submit():
        q.context = form.context.data
        q.answer = form.answer.data
        if qt == "选择题":
            q.options = form.options.data
        return render_template('ques/message.html', message='试题更新成功')
    db.session.add(q)
    db.session.commit()
    return render_template('ques/edit.html', form=form)


@ques.route("/delete/<int:qid>", methods=['GET', 'POST'])
def delete(qid):
    q = Question.query.filter_by(id=qid).first()
    db.session.delete(q)
    db.session.commit()
    return render_template('ques/message.html', message='试题删除成功')


@ques.route("/add/<subject>", methods=['GET', 'POST'])
def add(subject):
    form = AddQuestionForm()

    return render_template('ques/add.html', form=form, subject=subject)


@ques.route("/add/results/<qt_name>", methods=['GET', 'POST'])
def addResults(qt_name):
    kn_points = KnowledgePoint.query.all()
    form = AddQuestionForm()

    # qt = QuestionType.query.filter_by(name=qt_name).first()
    # if form.validate_on_submit():
    #     q.context = form.context.data
    #     print form.context.data
    #     print 'asdasd'
    #     print q.context
    #     if qt_name == '选择题':
    #         q.options = form.options.data
    #     q.answer = form.answer.data
    #     q.degree = form.degree.data
    #
    #     q.kn_point_counts = form.kn_point_counts.data
    #     q.kn_point_id = form.kn_point_id.data
    #     q.question_type_id = qt.id
    #     db.session.add(q)
    #     db.session.commit()
    #     return redirect(url_for('ques.addResults'), message="试题添加成功")
    return render_template('ques/add.html', form=form, qt_name=qt_name, kn_points=kn_points)


@ques.route("/add/complete/<qt_name>", methods=['GET', 'POST'])
def addComplete(qt_name):
    q = Question()
    qt = QuestionType.query.filter_by(name=qt_name).first()
    print qt.name
    form = AddQuestionForm()
    q.context = form.context.data
    subject = '数据结构'

    q.subject_id = Subject.query.filter_by(name=subject).first().id

    if qt_name == '选择题':
        q.options = form.options.data
    q.answer = form.answer.data
    q.degree = int(request.form['kn_point_degree'])

    q.kn_point_counts = int(request.form['kn_point_counts'])
    q.kn_point_id = KnowledgePoint.query.filter_by(name=request.form['kn_point_id_1']).first().id
    q.question_type_id = qt.id

    db.session.add(q)
    db.session.commit()
    qt_name = ''
    return render_template('ques/add.html', form=form, qt_name=qt_name, message='添加成功')


@ques.route("/organize", methods=['POST', 'GET'])
def organize():
    form = OrganizeQuestionForm()
    q_choice_list = Question.query.filter_by(question_type_id=QuestionType.query.filter_by(name="选择题").first().id).all()
    q_judgement_list = Question.query.filter_by(
        question_type_id=QuestionType.query.filter_by(name="判断题").first().id).all()
    q_blank_list = Question.query.filter_by(question_type_id=QuestionType.query.filter_by(name="填空题").first().id).all()

    ques_num = {
        '选择题个数': len(q_choice_list),
        '判断题个数': len(q_judgement_list),
        '填空题个数': len(q_blank_list),
    }
    if form.validate_on_submit():
        choice_num = int(form.choice_num.data)

        judgement_num = int(form.judgement_num.data)
        blank_num = int(form.blank_num.data)

        q_choice_list = Question.query.filter_by(
            question_type_id=QuestionType.query.filter_by(name="选择题").first().id).all()
        q_judgement_list = Question.query.filter_by(
            question_type_id=QuestionType.query.filter_by(name="判断题").first().id).all()
        q_blank_list = Question.query.filter_by(
            question_type_id=QuestionType.query.filter_by(name="填空题").first().id).all()

        results = organizeAlgorithm(choice_num, q_choice_list) + \
                  organizeAlgorithm(judgement_num, q_judgement_list) + \
                  organizeAlgorithm(blank_num, q_blank_list)
        ques_num = {
            '选择题个数': len(q_choice_list),
            '判断题个数': len(q_judgement_list),
            '填空题个数': len(q_blank_list),
        }

        filename = os.path.join('/home/zhuzw/git/flasky/downloads/paper',
                                datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '.doc')

        answername = os.path.join('/home/zhuzw/git/flasky/downloads/answer',
                                  datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '_answer' + '.doc')

        document = Document()
        document_answer = Document()
        document.add_paragraph(unicode('选择题：', 'utf-8'))
        document_answer.add_paragraph(unicode('选择题：', 'utf-8'))
        for q1 in organizeAlgorithm(choice_num, q_choice_list):
            document.add_paragraph(q1.context.decode('utf-8'))
            document.add_paragraph(q1.options.decode('utf-8'))
            document_answer.add_paragraph(q1.answer.decode('utf-8'))

        document.add_paragraph(unicode('判断题：', 'utf-8'))
        document_answer.add_paragraph(unicode('判断题：', 'utf-8'))
        for q2 in organizeAlgorithm(judgement_num, q_judgement_list):
            document.add_paragraph(q2.context.decode('utf-8'))
            document_answer.add_paragraph(q2.answer.decode('utf-8'))

        document.add_paragraph(unicode('填空题：', 'utf-8'))
        document_answer.add_paragraph(unicode('填空题：', 'utf-8'))
        for q3 in organizeAlgorithm(blank_num, q_blank_list):
            document.add_paragraph(q3.context.decode('utf-8'))
            document_answer.add_paragraph(q3.answer.decode('utf-8'))

        document.save(filename)
        document_answer.save(answername)

        return render_template("ques/organize.html",
                               filename=datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '.doc',
                               answer_name=datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '_answer' + '.doc',
                               results=results, count=len(results), form=form, ques_num=ques_num)

    return render_template("ques/organize.html", form=form, ques_num=ques_num)


@ques.route("/organize/results", methods=['POST', 'GET'])
def organizeResults():
    form = OrganizeQuestionForm()
    choice_num = int(form.choice_num.data)

    judgement_num = int(form.judgement_num.data)
    blank_num = int(form.blank_num.data)

    q_choice_list = Question.query.filter_by(question_type_id=QuestionType.query.filter_by(name="选择题").first().id).all()
    q_judgement_list = Question.query.filter_by(
        question_type_id=QuestionType.query.filter_by(name="判断题").first().id).all()
    q_blank_list = Question.query.filter_by(question_type_id=QuestionType.query.filter_by(name="填空题").first().id).all()

    results = organizeAlgorithm(choice_num, q_choice_list) + \
              organizeAlgorithm(judgement_num, q_judgement_list) + \
              organizeAlgorithm(blank_num, q_blank_list)
    ques_num = {
        '选择题个数': len(q_choice_list),
        '判断题个数': len(q_judgement_list),
        '填空题个数': len(q_blank_list),
    }

    filename = os.path.join('/home/zhuzw/git/flasky/downloads/paper',
                            datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '.doc')

    answername = os.path.join('/home/zhuzw/git/flasky/downloads/answer',
                              datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '_answer' + '.doc')

    document = Document()
    document_answer = Document()
    document.add_paragraph(unicode('选择题：', 'utf-8'))
    document_answer.add_paragraph(unicode('选择题：', 'utf-8'))
    for q1 in organizeAlgorithm(choice_num, q_choice_list):
        document.add_paragraph(q1.context.decode('utf-8'))
        document.add_paragraph(q1.options.decode('utf-8'))
        document_answer.add_paragraph(q1.answer.decode('utf-8'))

    document.add_paragraph(unicode('判断题：', 'utf-8'))
    document_answer.add_paragraph(unicode('判断题：', 'utf-8'))
    for q2 in organizeAlgorithm(judgement_num, q_judgement_list):
        document.add_paragraph(q2.context.decode('utf-8'))
        document_answer.add_paragraph(q2.answer.decode('utf-8'))

    document.add_paragraph(unicode('填空题：', 'utf-8'))
    document_answer.add_paragraph(unicode('填空题：', 'utf-8'))
    for q3 in organizeAlgorithm(blank_num, q_blank_list):
        document.add_paragraph(q3.context.decode('utf-8'))
        document_answer.add_paragraph(q3.answer.decode('utf-8'))

    document.save(filename)
    document_answer.save(answername)

    return render_template("ques/organize.html",
                           filename=datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '.doc',
                           answer_name=datetime.datetime.now().strftime("%Y%m%d%H%M%S") + '_answer' + '.doc',
                           results=results, count=len(results), form=form, ques_num=ques_num)


# 递归 防止重复，以后开发
def organizeAlgorithm(num, question_list):
    r = []

    for i in range(0, num):
        tmp = question_list[random.randint(0, len(question_list) - 1)]
        r.append(tmp)
    return r


# 为了测试，预定义1到选择题
@ques.route("/examOnline", methods=['POST', 'GET'])
def examOnline():
    form = AnswerForm()
    q_choice_list = Question.query.filter_by(question_type_id=QuestionType.query.filter_by(name="选择题").first().id).all()
    # q_judgement_list = Question.query.filter_by(question_type_id=QuestionType.query.filter_by(name="判断题").first().id).all()
    # q_blank_list = Question.query.filter_by(question_type_id=QuestionType.query.filter_by(name="填空题").first().id).all()

    results = {
        "选择题": organizeAlgorithm(1, q_choice_list),
        # "判断题": organizeAlgorithm(1, q_judgement_list),
        # "填空题": organizeAlgorithm(1, q_blank_list),
    }

    return render_template("ques/exam_online.html", results=results['选择题'][0], form=form)


@ques.route("/checkAnswer/<int:qid>", methods=['POST', 'GET'])
def checkAnswer(qid):
    form = AnswerForm()
    answer = form.choiceAnswer.data

    q = Question.query.filter_by(id=qid).first()
    # qt = QuestionType.query.filter_by(name="选择题").first()
    q.appearance += 1
    if answer == q.answer:
        message = "回答正确"
        q.accuracy = 1.0 * (q.accuracy * (q.appearance - 1) + 1) / q.appearance
    else:
        message = "回答错误"
        q.accuracy = 1.0 * (q.accuracy * (q.appearance - 1)) / q.appearance
    db.session.add(q)
    db.session.commit()

    return render_template('ques/score.html', message=message, q=q)


@ques.route("/addSubject", methods=['POST', 'GET'])
def addSubject():
    form = SubjectForm()
    message = ''
    if form.validate_on_submit():
        name = form.name.data
        db.session.add(Subject(name=name))
        db.session.commit()
        message = '添加成功'

    return render_template('ques/addSubject.html', form=form, message=message)


@ques.route("/addChapter", methods=['POST', 'GET'])
def addChapter():
    form = ChapterForm()

    subject_list = Subject.query.all()
    if form.validate_on_submit():
        subject_id = request.form['subject_id']
        form = ChapterForm()
        name = form.name.data
        db.session.add(Chapter(name=name, subject_id=subject_id))
        db.session.commit()
        message = '添加成功'
        return render_template('ques/message.html', message=message)

    return render_template('ques/addChapter.html', form=form, subject_list=subject_list)


@ques.route("/addChapterComplete", methods=['POST', 'GET'])
def addChapterComplete():
    subject_id = request.form['subject_id']
    form = ChapterForm()
    name = form.name.data
    db.session.add(Chapter(name=name, subject_id=subject_id))
    db.session.commit()
    message = '添加成功'
    return render_template('ques/message.html', message=message)


@ques.route("/addKnPoints", methods=['POST', 'GET'])
def addKnPoints():
    form = ChapterForm()
    subject_list = Subject.query.all()
    chapter_list = Chapter.query.all()

    if form.validate_on_submit():
        subject_id = request.form['subject_id']
        knpoint_id = request.form['knpoint_id']
        form = ChapterForm()
        name = form.name.data
        db.session.add(KnowledgePoint(name=name, subject_id=subject_id, knpoint_id=knpoint_id))
        db.session.commit()
        message = '添加成功'
        return render_template('ques/message.html', message=message)

    return render_template('ques/addKnpoints.html', form=form, subject_list=subject_list, chapter_list=chapter_list)


@ques.route("/addKnPointsComplete", methods=['POST', 'GET'])
def addKnPointsComplete():
    subject_id = request.form['subject_id']
    knpoint_id = request.form['knpoint_id']
    form = ChapterForm()
    name = form.name.data
    db.session.add(KnowledgePoint(name=name, subject_id=subject_id, knpoint_id=knpoint_id))
    db.session.commit()
    message = '添加成功'
    return render_template('ques/message.html', message=message)


def open_excel(filename):
    try:
        data = xlrd.open_workbook(filename)
        return data
    except Exception, e:
        print str(e)


# 根据索引获取Excel表格中的数据   参数:file：Excel文件路径     colnameindex：表头列名所在行的所以  ，by_index：表的索引
def excel_table_byindex(filename, colnameindex=0, by_index=0):
    data = open_excel(filename)
    table = data.sheets()[by_index]
    nrows = table.nrows  # 行数
    ncols = table.ncols  # 列数
    colnames = table.row_values(colnameindex)  # 某一行数据
    list = []
    for rownum in range(1, nrows):

        row = table.row_values(rownum)
        if row:
            app = {}
            for i in range(len(colnames)):
                app[colnames[i]] = row[i]
            list.append(app)
    return list


UPLOAD_FOLDER = '/home/zhuzw/git/flasky/uploads'


@ques.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'GET':
        return render_template('ques/upload.html')
    elif request.method == 'POST':
        f = request.files['file']

        fname = secure_filename(f.filename)

        filename = os.path.join(UPLOAD_FOLDER, fname)
        f.save(filename)

        print 'path:', filename
        datas = excel_table_byindex(filename)

        for row in datas:
            ques = Question()
            ques.context = row['context']
            ques.answer = row['answer']
            ques.degree = row['degree']

            if row['ques_type'] == '选择题':
                ques.options = row['options']

            ques.subject_id = Subject.query.filter_by(name=row['subject']).first().id
            ques.kn_point_id = KnowledgePoint.query.filter_by(name=row['kn_points']).first().id
            ques.question_type_id = QuestionType.query.filter_by(name=row['ques_type']).first().id
            ques.professor_id = User.query.filter_by(username='prozhu').first().id

            db.session.add(ques)
        db.session.commit()

        return render_template('ques/message.html', message='上传成功')


@ques.route('/download_paper/<filename>', methods=['GET', 'POST'])
def download_paper(filename):
    # strIO = StringIO.StringIO()
    # # strIO.write('Hello from Dan Jacob and Stephane Wirtel !')
    # strIO.seek(0)
    # return send_file(strIO, attachment_filename=filename, as_attachment=True)
    return send_from_directory('/home/zhuzw/git/flasky/downloads/paper', filename, as_attachment=True)


@ques.route('/download_answer/<answer_name>', methods=['GET', 'POST'])
def download_answer(answer_name):
    # strIO = StringIO.StringIO()
    # # strIO.write('Hello from Dan Jacob and Stephane Wirtel !')
    # strIO.seek(0)
    # return send_file(strIO, attachment_filename=filename, as_attachment=True)
    return send_from_directory('/home/zhuzw/git/flasky/downloads/answer', answer_name, as_attachment=True)


@ques.route('/all_subject', methods=['GET', 'POST'])
def all_subject():
    subjects = Subject.query.all()
    # results = []
    # for s in subjects:
    #     results.append(s)
    # return jsonify(results)
    return subjects